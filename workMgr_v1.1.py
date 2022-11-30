import sys
import tkinter as tk
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
import sqlite3
import time
import os
import configparser as configparser
from tkinter import filedialog
from tkinter.filedialog import askdirectory
import docx
from docx import Document
from tkinter.ttk import Separator

# 数据库初始化
conn = sqlite3.connect('workspace.db')
cur = conn.cursor()

#读取配置文件
cong=configparser.ConfigParser()
path='config.ini'
if os.path.exists(path):
	cong.read(path,'utf-8-sig')
	monthRptName=cong.get('monthrpt','月报保存名称')
else:
	print('>>>配置文件不存在!')
	input('--按回车键继续--')
	sys.exit()


def getPathConfig():
    cmd_getthepat="select * from paths"
    cur.execute(cmd_getthepat)
    res_t=cur.fetchall()
    #print("getPathConfig res_t：",res_t)
    if not res_t==[]:
        WeekModelPath.set(res_t[0][1])
        MonthModelPath.set(res_t[0][2])
        WeekSavePath.set(res_t[0][3])
        MonthSavePath.set(res_t[0][4])
    # else:
    #     WeekModelPath.set('')
    #     MonthModelPath.set('')
    #     WeekSavePath.set('')
    #     MonthSavePath.set('')

def initTable():
    try:
        creat_table_1 = '''CREATE TABLE scores
           (id INTEGER PRIMARY KEY,
           日期 date,
           星期几 TEXT,
           稿件名称 TEXT,
            工作类型 TEXT,
            加工页数 NUMBER,
            加工字数 NUMBER,
            系数 NUMBER,
            周报标签 NUMBER,
            月报标签 NUMBER
            );'''
        # init_table_2="INSERT INTO scores VALUES(NULL,(?),0)"
        cur.execute(creat_table_1)
    except:
        print("数据表scores已存在")
    try:
        creat_table_2 = '''CREATE TABLE month_scores
                   (id INTEGER PRIMARY KEY,
                   日期 date,
                    工作量 NUMBER
                    );'''
        cur.execute(creat_table_2)
    except:
        print("数据表month_scores已存在")
    try:
        creat_table_3 = '''CREATE TABLE paths
                   (id INTEGER PRIMARY KEY,
                   WeekModelPath TEXT,
                   MonthModelPath TEXT,
                   WeekSavePath TEXT,
                   MonthSavePath TEXT
                    );'''
        cur.execute(creat_table_3)
        conn.commit()
        getPathConfig()
    except:
        getPathConfig()
        print("数据表paths已存在")
    try:
        creat_table_4 = '''CREATE TABLE book_scores
           (id INTEGER PRIMARY KEY,
           稿件名称 TEXT,
            总页数 NUMBER,
            总字数 NUMBER,
            说明 TEXT,
            寄回标签 NUMBER
            );'''
        # init_table_2="INSERT INTO scores VALUES(NULL,(?),0)"
        cur.execute(creat_table_4)
    except:
        print("数据表book_scores已存在")



def getAllBookName():
    res=[]
    cmd= "select * from book_scores"
    cur.execute(cmd)
    res_t = cur.fetchall()
    #print("getAllBookName res_t：", res_t)
    if not res_t == []:
        for item in res_t:
            res.append(item[1])
    res.reverse()
    return  res
# ------------------窗体----------------------

window = tk.Tk()
window.title('云海-工作管理系统')

# 设置窗口大小变量
width = 1000
height = 800
# 窗口居中，获取屏幕尺寸以计算布局参数，使窗口居屏幕中央
screenwidth = window.winfo_screenwidth()
screenheight = window.winfo_screenheight()
size_geo = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
window.geometry(size_geo)
label_title_1 = tk.Label(window, text="添加工作记录",
                         font=(10,),
                         padx=1, pady=1,
                         borderwidth=2,relief='raised').grid(row=0, column=0,ipadx=10,ipady=10)

str_show_tmp = tk.StringVar()
str_show_tmp.set("测试信息展示区")
# 添加信息区域
frame1 = tk.Frame(window)
frame1.grid(row=1)
# 展示信息区域
frame2 = tk.Frame(window)
frame2.grid(row=3)
# 功能区，生成周报，月报等
frame3 = tk.Frame(window)
frame3.grid(row=2)
#稿件基础信息
label_book_title = tk.Label(window, text="添加稿件基础信息",font=(10,),padx=1, pady=1,borderwidth=2,relief='raised').grid(row=4,ipadx=10,ipady=10)
frame4 = tk.Frame(window)
frame4.grid(row=5)
frame5 = tk.Frame(window)
frame5.grid(row=6)
# 工作记录添加区控件
label_name = tk.Label(frame1, text="稿件名称：").grid(row=1, column=1)
label_type = tk.Label(frame1, text="工作类型：").grid(row=1, column=3)
label_page = tk.Label(frame1, text="加工页数：").grid(row=1, column=5)
label_pagenum = tk.Label(frame1, text="页").grid(row=1, column=7)
#label_word = tk.Label(frame1, text="加工字数：").grid(row=1, column=9)
#label_wordnum = tk.Label(frame1, text="万").grid(row=1, column=11)
label_xishu = tk.Label(frame1, text="系数(默认为1):").grid(row=1, column=12)
#label_xishu_info = tk.Label(frame1, text="*系数默认为1:").grid(row=2, column=12)

name_var = tk.StringVar()
type_var = tk.StringVar()
pagenum_var = tk.StringVar()
wordnum_var = tk.StringVar()
xishu_var = tk.StringVar()
#entry_name = tk.Entry(frame1, width=35,textvariable=name_var)
entry_name=ttk.Combobox(frame1,width=35, textvariable=name_var)

entry_name.grid(row=1, column=2, padx=5)
box_type = ttk.Combobox(frame1,width=10, textvariable=type_var)
box_type['value'] = ('第一遍加工', '第二遍加工', '原稿寄回', '读校样')
box_type.grid(row=1, column=4)
box_type.configure(state="readonly")
entry_pagenum = tk.Entry(frame1, width=10, textvariable=pagenum_var)
entry_pagenum.grid(row=1, column=6)
#entry_wordnum = tk.Entry(frame1, width=10, textvariable=wordnum_var)
#entry_wordnum.grid(row=1, column=10)
entry_xishu = tk.Entry(frame1, width=10, textvariable=xishu_var)
entry_xishu.grid(row=1, column=13)
xishu_var.set(1)
#工作记录表格
yscroll = Scrollbar(frame2, orient=VERTICAL, width=30)
tree = ttk.Treeview(frame2, show='headings', yscrollcommand=yscroll.set)  # #创建表格对象
tree["columns"] = (
"日期", "星期几", "稿件名称", "工作类型", "加工页数", "加工字数", "系数")  # #定义列
tree.column("日期", width=80)
tree.column("星期几", width=80)
tree.column("稿件名称", width=300)
tree.column("工作类型", width=80)  # #设置列
tree.column("加工页数", width=80)
tree.column("加工字数", width=80)
tree.column("系数", width=50)
tree.heading("日期", text="日期")
tree.heading("星期几", text="Weekday")
tree.heading("稿件名称", text="稿件名称")  # #设置显示的表头名
tree.heading("工作类型", text="工作类型")
tree.heading("加工页数", text="加工页数")
tree.heading("加工字数", text="加工字数/万")
tree.heading("系数", text="系数")
yscroll.config(command=tree.yview)
yscroll.grid(row=0, column=1, sticky=S + W + E + N)
tree.grid(row=0, column=0)
label_monthwork_text = tk.StringVar()
label_monthwork = tk.Label(frame2, textvariable=label_monthwork_text).grid(row=1)

#--------------稿件基础信息区--------------

label_book_name = tk.Label(frame4, text="稿件名称：").grid(row=1, column=0)
book_name_var = tk.StringVar()
entry_book_name = tk.Entry(frame4, width=35,textvariable=book_name_var)
entry_book_name.grid(row=1, column=1)
label_book_pagenum = tk.Label(frame4, text="总页数：").grid(row=1, column=2)
book_pagenum_var = tk.StringVar()
entry_book_pagenum = tk.Entry(frame4, width=8,textvariable=book_pagenum_var)
entry_book_pagenum.grid(row=1, column=3)
label_book_wordnum = tk.Label(frame4, text="总字数：").grid(row=1, column=4)
book_wordnum_var = tk.StringVar()
entry_book_wordnum = tk.Entry(frame4, width=8,textvariable=book_wordnum_var)
entry_book_wordnum.grid(row=1, column=5)
label_book_wordnum_unit = tk.Label(frame4, text="万").grid(row=1, column=6)
label_book_info = tk.Label(frame4, text="说明：").grid(row=1, column=7)
book_info_var = tk.StringVar()
entry_book_info = tk.Entry(frame4, width=15,textvariable=book_info_var)
entry_book_info.grid(row=1, column=8)
#稿件信息表格
yscroll_book = Scrollbar(frame5, orient=VERTICAL, width=30)
tree_book = ttk.Treeview(frame5, show='headings', yscrollcommand=yscroll_book.set)  # #创建表格对象
tree_book["columns"] = (
"稿件名称", "总页数", "总字数", "说明")  # #定义列
tree_book.column("稿件名称", width=300) # #设置列
tree_book.column("总页数", width=75)
tree_book.column("总字数", width=75)
tree_book.column("说明", width=300)
tree_book.heading("稿件名称", text="稿件名称")  # #设置显示的表头名
tree_book.heading("总页数", text="总页数")
tree_book.heading("总字数", text="总字数/万")
tree_book.heading("说明", text="说明")
yscroll_book.config(command=tree_book.yview)
yscroll_book.grid(row=0, column=1, sticky=S + W + E + N)
tree_book.grid(row=0, column=0,pady=10)
label_bookbox_text = tk.StringVar()
label_bookbox = tk.Label(frame5, textvariable=label_bookbox_text).grid(row=1)
label_bookbox_text.set("*绿色为原稿寄回标记")
#----------------全局----------------------
WeekModelPath=tk.StringVar()
MonthModelPath=tk.StringVar()
WeekSavePath=tk.StringVar()
MonthSavePath=tk.StringVar()
# -----------------功能--------------------
initTable()
entry_name['value']=getAllBookName()
#刷新工作记录表格
def refresh():
    for item in tree.get_children():
        tree.delete(item)
    select_all_cmd = "SELECT * FROM scores"
    cur.execute(select_all_cmd)
    result_list = cur.fetchall()
    # print("reslut_list:",result_list)
    data = []
    for i in range(0, len(result_list)):
        for j in range(1, len(result_list[i])):
            data.append(result_list[i][j])
        # print("data；")
        # print(data)
        if (data[len(data) - 1] == 0) & (data[len(data) - 2] == 0):
            tree.insert('', 0, values=data)
        else:
            if (data[len(data) - 1] == 1) & (data[len(data) - 2] == 1):
                tree.insert('', 0, values=data, tags='month&week')
            else:
                if data[len(data) - 1] == 1:
                    tree.insert('', 0, values=data, tags='month')
                if data[len(data) - 2] == 1:
                    tree.insert('', 0, values=data, tags='week')
        data.clear()
    tree.tag_configure('month', background='gray')
    tree.tag_configure('week', background='orange')
    tree.tag_configure('month&week', background='red')
    # 信息展示
    # label_list = []
    # for i in range(1, 11):
    #     label_tmp = tk.Label(frame2)
    #     label_tmp.grid(row=i - 1, pady=10)
    #     label_list.append(label_tmp)
    # label_var_list = []
    # for i in range(1, 11):
    #     var_tmp = tk.StringVar()
    #     var_tmp.set(str(i - 1))
    #     label_var_list.append(var_tmp)
    # for i in range(1, 11):
    #     # print(label_list[i-1])
    #     label_list[i - 1]['textvariable'] = label_var_list[i - 1]
#刷新稿件基础信息表格
def refreshbookbox():
    for item in tree_book.get_children():
        tree_book.delete(item)
    select_all_cmd = "SELECT * FROM book_scores"
    cur.execute(select_all_cmd)
    result_list = cur.fetchall()
    # print("reslut_list:",result_list)
    data = []
    for i in range(0, len(result_list)):
        for j in range(1, len(result_list[i])):
            data.append(result_list[i][j])
        # print("data；")
        # print(data)
        if (data[len(data) - 1] == 0):
            tree_book.insert('', 0, values=data)
        elif (data[len(data) - 1] == 1):
            tree_book.insert('', 0, values=data, tags='finished')
        data.clear()
    tree_book.tag_configure('finished', background='green')
    #tree_book.tag_configure('week', background='orange')
    #tree_book.tag_configure('month&week', background='red')
    entry_name['value']=getAllBookName()

# 判断是否为合法系数
def isxishu(param):
    if param.isdigit():
        if int(param) < 10:
            #print(int(param))
            return True
    else:
        if param.find(".") != -1:
            p1 = param.split(".")[0]
            p2 = param.split(".")[1]
            if p1.isdigit() & p2.isdigit() & (int(p1) < 10) & (int(p2) < 10):
                return True
            else:
                return False
        else:
            return False

# 添加信息区控件功能


def existBook(param):
    cmd="select * from book_scores where 稿件名称=(?)"
    cur.execute(cmd,(param,))
    conn.commit()
    res=cur.fetchall()
    if res == []:
        return False
    else: return True


def getCurrentWorkWordnum(pagenum, bookname):
    cmd="select 总页数,总字数 from book_scores where 稿件名称=(?)"
    cur.execute(cmd,(bookname,))
    conn.commit()
    res_t=cur.fetchall()
    #print("current res_t:",res_t)
    book_pagenum=res_t[0][0]
    book_wordnum=res_t[0][1]
    res=(int(pagenum)/book_pagenum)*book_wordnum
    res="{0:.2f}".format(res)
    return res


def updateBookState(bookname):
    cmd="update book_scores set 寄回标签=1 where 稿件名称=(?)"
    cur.execute(cmd,(bookname,))
    conn.commit()
    return


def additem():
    if not existBook(str(entry_name.get())):
        messagebox.showerror("找不到稿件", "请先添加该稿件基础信息")
        book_name_var.set(entry_name.get())
    elif (entry_pagenum.get().isdigit() or entry_pagenum.get() == '') & isxishu(
            entry_xishu.get()):
        print("添加了新项")
        t = time.localtime()
        today = time.strftime("%Y-%m-%d", t)
        weekday_num = int(time.strftime("%w", t))
        weekdays = ('周日', '周一', '周二', '周三', '周四', '周五', '周六')
        weekday = weekdays[weekday_num]
        # print(today)
        # print(weekday)
        if entry_pagenum.get() == '':
            pagenum = 0
        else:
            pagenum = entry_pagenum.get()
        current_wordnum=getCurrentWorkWordnum(pagenum,entry_name.get()) #根据页数与总字数计算本次加工字数
        if box_type.get()=="原稿寄回":
            updateBookState(entry_name.get())
        add_item = "INSERT INTO scores VALUES(NULL,'" + today + "','" + weekday + "','" + entry_name.get() + "','" + box_type.get() + "'," + str(
            pagenum) + "," + current_wordnum + "," + entry_xishu.get() + ",0,0); "
        cur.execute(add_item)
        conn.commit()
        name_var.set('')
        type_var.set('')
        wordnum_var.set('')
        xishu_var.set('')
        pagenum_var.set('')
        refresh()
        refreshbookbox()
        # str_show_tmp.set("稿件名称："+entry_name.get()+"，加工类型："+box_type.get()+"，加工页数："+entry_pagenum.get()+"页"+"加工字数："+entry_wordnum.get()+"万字"+" 系数："+entry_xishu.get())
    else:
        messagebox.showerror("错误", "请正确填写所有信息")
        if not entry_pagenum.get().isdigit(): pagenum_var.set('')
        if not isxishu(entry_xishu.get()): xishu_var.set('')
    return

# 功能区控件功能

def deleteAll():
    delete_all_cmd = "DELETE FROM scores"
    cur.execute(delete_all_cmd)
    conn.commit()
    refresh()


def deleteAllorNot():
    # 弹出对话框
    result = tk.messagebox.askokcancel(title='', message='删除全部记录？')
    if result:
        deleteAll()

# 获取最近一次记录的id
def getmaxid_in_score():
    find_maxid_cmd = "select max(id) from scores;"
    cur.execute(find_maxid_cmd)
    maxid = cur.fetchall()[0]
    # print("maxid:", maxid[0])
    if maxid[0] == None:
        return [0, ]
    else:
        return maxid

def deleteRecent():
    maxid = getmaxid_in_score()
    # print("maxid:")
    # print(maxid)
    cmd_1="select 稿件名称,工作类型 from scores where id=(?)"
    cur.execute(cmd_1,maxid)
    conn.commit()
    res_t=cur.fetchall()
    if res_t[0][1]=="原稿寄回":
        cmd_2 = "update book_scores set 寄回标签=0 where 稿件名称=(?)"
        cur.execute(cmd_2,(res_t[0][0],))
        conn.commit()
    delete_recent_cmd = "delete from scores where id=(?)"
    cur.execute(delete_recent_cmd, maxid)
    conn.commit()
    refresh()
    refreshbookbox()

def deleteRecorNot():
    result = tk.messagebox.askokcancel(title='', message='删除上一条记录？')
    if result:
        deleteRecent()

def getworkdata(worktype, last_tag_id, maxid):
    cmd_get = "select 稿件名称,工作类型,sum(加工页数),sum(加工字数),系数 from (select * from scores where 工作类型=(?) and id>(?) and id<=(?)) group by 稿件名称"
    cur.execute(cmd_get, (worktype, last_tag_id, maxid))
    result = cur.fetchall()
    return result

def getworkDatas(counttype):
    workdatas = []
    # 寻找上一个周报tag,按照id递减排序，
    # print("---------生成"+counttype+"数据-----------：")
    cmd_1 = "select * from scores where " + counttype + "标签=1 order by id desc limit 0,1"
    cur.execute(cmd_1)
    c_result = cur.fetchone()
    # print("c_result:",c_result)
    if c_result == None:
        c_result_id = 0
    else:
        c_result_id = c_result[0]
    # print("c_result_id:",c_result_id)
    maxid = getmaxid_in_score()
    # print("maxid2=",maxid)
    if c_result_id < maxid[0]:
        if counttype == "周报":
            # 获取周报数据
            workdata_1 = getworkdata('第一遍加工', c_result_id, maxid[0])
            workdata_2 = getworkdata('第二遍加工', c_result_id, maxid[0])
            workdata_3 = getworkdata('读校样', c_result_id, maxid[0])
        elif counttype == "月报":
            workdata_1 = getworkdata('第一遍加工', c_result_id, maxid[0])
            workdata_2 = getworkdata('原稿寄回', c_result_id, maxid[0])
            workdata_3 = getworkdata('读校样', c_result_id, maxid[0])
        workdatas.extend(workdata_1)
        workdatas.extend(workdata_2)
        workdatas.extend(workdata_3)
        #print("getWorkDatas:", workdatas)
    return workdatas

def exAndSave_MonthRpt(modelpath, savepath, extext):
    doc = Document(modelpath)
    tables = doc.tables
    t = time.localtime()
    today = time.strftime("%Y%m%d", t)
    month=time.strftime("%m", t)
    savepath=savepath+"/"+monthRptName+".docx"
    if len(tables) != 0:
        print('分析表格')
        tb1 = tables[0]
        rows = tb1.rows
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text
                        if not text.find("date")==-1:
                            run.text=run.text.replace("date",today)
                            print(run.text)
                            continue
                        elif not text.find("m")==-1:
                            run.text=run.text.replace("m",month)
                            print(cell.text)
                            continue
                        elif not text.find('{{text1}}')==-1:
                            run.text=run.text.replace('{{text1}}',extext[0])
                            print(run.text)
                            continue
                        elif not text.find('{{text2}}')==-1:
                            run.text=run.text.replace('{{text2}}',extext[1])
                            print(cell.text)
                            continue
                        elif not text.find('{{text3}}')==-1:
                            run.text=run.text.replace('{{text3}}',extext[2])
                            continue
                        # print(text)
        doc.save(savepath)
    else:
        # print('无表格，保存为'+savepath)
        doc.save(savepath)
    return

def exAndSave_WeekRpt(modelpath, savepath, extext):
    doc = Document(modelpath)
    t = time.localtime()
    today = time.strftime("%Y年%m月%d日", t)
    savepath=savepath+"/"+today+".txt"
    new_file = open(savepath, 'w')
    print(savepath)
    for p in doc.paragraphs:
        if '{{text1}}' in p.text:
            p.text = p.text.replace('{{text1}}', extext[0])
        new_file.write(p.text + '\r')
    new_file.close()
    return


def gentext(type,data):
    text_res_1= ""
    text_res_2 = ""
    if type=="月报":
        for i in range(0, len(data)):
            if data[i][1] == '第一遍加工':
                text_res_1= text_res_1 + "稿件名称 《" + data[i][0] + "》 字数 " + str(data[i][3]) + "万字（第一遍）\n"
            elif data[i][1] == '原稿寄回':
                text_res_1= text_res_1 + "稿件名称 《" + data[i][0] + "》 字数 " + str(data[i][3]) + "万字（已寄回）\n"
            elif data[i][1] == '读校样':
                text_res_2= text_res_2 + "稿件名称 《" + data[i][0] + "》 字数 " + str(data[i][3]) + "万字\n"
    elif type=="周报":
        for i in range(0, len(data)):
            if data[i][1] == '第一遍加工':
                text_res_1= text_res_1 + "加工稿件 《" + data[i][0] + "》 第一遍" + str(data[i][2]) + "页。"
            elif data[i][1] == '第二遍加工':
                text_res_1= text_res_1 + "加工稿件 《" + data[i][0] + "》 第二遍 " + str(data[i][2]) + "页。"
            elif data[i][1] == '读校样':
                text_res_2= text_res_2 + "读校样 《" + data[i][0] + "》 " + str(data[i][2]) + "页。"
    return [text_res_1,text_res_2]




def weekReport(data):
    #print("周报:", data)
    texts = gentext("周报", data)
    text_3 = genComment("周报")
    texts.append(text_3)
    print("周报:",texts)
    # 保存至文档
    exAndSave_WeekRpt(WeekModelPath.get(), WeekSavePath.get(), texts)


def genComment(type):
    comment = ''#comment内容待定
    if type=="周报":
        pass
    elif type=="月报":
        pass
    return comment

def monthReport(data):
    print("月报:",data)
    texts= gentext("月报", data)
    text_3=genComment("月报")
    texts.append(text_3)
    print("月报：",texts)
    #保存至文档
    exAndSave_MonthRpt(MonthModelPath.get(),MonthSavePath.get(),texts)

def genRpt(data, param):
    # 文件写操作
    if param == "周报":
        weekReport(data)
    if param == "月报":
        monthReport(data)
    print(">>>>>>生成了" + param)

def updatetag(counttype):
    # 为最后新记录加周报tag
    maxid = getmaxid_in_score()
    update_cmd = "update scores set " + counttype + "标签=1 where id=(?)"
    cur.execute(update_cmd, maxid)
    conn.commit()
    refresh()

def update_monthwork_inDB(today, param, bool):
    if not bool:  # False,插入
        add_item = "INSERT INTO month_scores VALUES(NULL,'" + today + "'," + param + ");"
    else:
        add_item = "update month_scores set 工作量=" + param + " where 日期='" + today + "'"
    cur.execute(add_item)
    conn.commit()

def genmonthMap():
    cmd_1 = "select * from month_scores"
    cur.execute(cmd_1)
    result_cmd_1 = cur.fetchall()
    print("result_cmd_1:", result_cmd_1)
    data = []
    for i in range(0, len(result_cmd_1)):
        for j in range(1, len(result_cmd_1[i])):
            data.append(result_cmd_1[i][j])
    print("monthMap；", data)

def updata_monthwork():
    # 在窗口中显示本月已完成的工作量
    workdatas = getworkDatas("月报")
    work_065 = 0
    work_02 = 0
    work_015 = 0
    for i in range(0, len(workdatas)):
        if workdatas[i][1] == '第一遍加工':
            work_065 = work_065 + workdatas[i][3]
        if workdatas[i][1] == "原稿寄回": work_02 = work_02 + workdatas[i][3]
        if workdatas[i][1] == "读校样": work_015 = work_015 + workdatas[i][3]
    monthwork_var = round(work_065 * 0.65 + work_02 * 0.2 + work_015 * 0.15, 2)
    label_monthwork_text.set("*橙色为周报标记，灰色为月报标记，红色为当天同时提交周报月报。本月目前工作量：" + str(monthwork_var) + "万字")
    return str(monthwork_var)

def write_month_into_DB(monthwork_var):
    t = time.localtime()
    today = time.strftime("%Y-%m", t)
    cmd_search = "select * from month_scores where 日期='" + today + "'"
    cur.execute(cmd_search)
    cmd_search_result = cur.fetchall()
    print("cmd_search_result:", cmd_search_result)
    if cmd_search_result == []:
        update_monthwork_inDB(today, str(monthwork_var), False)
    else:
        update_monthwork_inDB(today, str(monthwork_var), True)

def genmonthRpt():
    print("MonthModelPath:",MonthModelPath.get())
    if MonthModelPath.get()==''or MonthSavePath.get()=='':
        tk.messagebox.showinfo(title='信息', message="请先设置模板及保存位置及名称")
    else:
        result = tk.messagebox.askokcancel(title='', message='生成月报至 '+MonthSavePath.get()+'，此操作将添加月报标签')
        if result:
            monthdata = getworkDatas("月报")
            write_month_into_DB(updata_monthwork())
            genmonthMap()
            genRpt(monthdata, "月报")
            updatetag("月报")
            tk.messagebox.showinfo(title='信息', message="月报生成成功！")

def genweekRpt():
    if WeekModelPath.get()==''or WeekSavePath.get()=='':
        tk.messagebox.showinfo(title='信息', message="请先设置模板及保存位置")
    else:
        result = tk.messagebox.askokcancel(title='', message="生成周报至 "+WeekSavePath.get()+"，此操作将添加周报标签")
        if result:
            weekdata = getworkDatas("周报")
            updatetag("周报")
            genRpt(weekdata, "周报")
            tk.messagebox.showinfo(title='信息', message="周报生成成功！")

def genweekRpt_undo(type):
    maxid = getmaxid_in_score()
    update_cmd = "update scores set " + type + "标签=0 where id=(?)"
    cur.execute(update_cmd, maxid)
    conn.commit()
    refresh()


def selectWeekModelPath():
    settingWindow.wm_attributes("-topmost", 0)
    WeekModelPath.set(filedialog.askopenfilename())
    settingWindow.wm_attributes("-topmost", 1)

def selectMonthModelPath():
    settingWindow.wm_attributes("-topmost", 0)
    MonthModelPath.set(filedialog.askopenfilename())
    settingWindow.wm_attributes("-topmost", 1)


def selectWeekSavePath():
    settingWindow.wm_attributes("-topmost", 0)
    WeekSavePath.set(askdirectory())
    settingWindow.wm_attributes("-topmost", 1)

def selectMonthSavePath():
    settingWindow.wm_attributes("-topmost", 0)
    MonthSavePath.set(askdirectory())
    settingWindow.wm_attributes("-topmost", 1)

def applySetting():
    cmd_savethepath_check="select * from paths"
    cur.execute(cmd_savethepath_check)
    res_t=cur.fetchall()
    if res_t==[]:
        cmd_savethepath = "INSERT INTO paths VALUES(NULL,'" + WeekModelPath.get() + "','" + MonthModelPath.get() + "','"+WeekSavePath.get()+"','"+MonthSavePath.get()+"');"
    else:
        cmd_savethepath = "update paths set WeekModelPath='" + WeekModelPath.get() + "',MonthModelPath='"+MonthModelPath.get()+"', WeekSavePath='"+WeekSavePath.get()+"',MonthSavePath='"+MonthSavePath.get()+"' where id=1"
    cur.execute(cmd_savethepath)
    conn.commit()
    window.attributes("-disabled", 0)
    settingWindow.destroy()


def setmonthRptName():
    os.system("notepad config.ini")
    return


def creatsettingWindow():
    window.attributes("-disabled", 1)
    global settingWindow
    settingWindow = tk.Toplevel(window)
    swidth=800
    sheight=320
    ssize_geo = '%dx%d+%d+%d' % (swidth, sheight, (screenwidth - swidth) / 2, (screenheight - sheight) / 2)
    settingWindow.geometry(ssize_geo)
    settingWindow.geometry('800x320')
    settingWindow.title('设置')
    settingWindow.wm_attributes("-topmost", 1)
    but_setWeekModelPath = tk.Button(settingWindow, text="选择周报模板", command=lambda: [selectWeekModelPath()]).grid(
        row=0,column=0,  padx=5, pady=10)
    label_WeekModelPath= tk.Label(settingWindow, textvariable=WeekModelPath).grid(row=0, column=1)
    but_setMonthModelPath = tk.Button(settingWindow, text="选择月报模板", command=lambda: [selectMonthModelPath()]).grid(
        row=1, column=0, padx=5, pady=10)
    label_WeekMonthPath = tk.Label(settingWindow, textvariable=MonthModelPath).grid(row=1, column=1)
    but_setWeekSavePath = tk.Button(settingWindow, text="周报保存路径", command=lambda: [selectWeekSavePath()]).grid(
        row=2, column=0, padx=5, pady=10)
    label_WeekSavePath = tk.Label(settingWindow, textvariable=WeekSavePath).grid(row=2, column=1)
    but_setMonthSavePath = tk.Button(settingWindow, text="月报保存路径", command=lambda: [selectMonthSavePath()]).grid(
        row=3, column=0, padx=5, pady=10)
    but_setMonthSaveName = tk.Button(settingWindow, text="修改月报名称\n(修改后请重启软件)", command=lambda: [setmonthRptName()]).grid(
        row=4, column=0, padx=5, pady=10)
    label_WeekSavePath = tk.Label(settingWindow, textvariable=MonthSavePath).grid(row=3, column=1)
    but_makesure=tk.Button(settingWindow, text="确定",command=lambda :applySetting()).grid(row=4, column=2,ipadx=8,ipady=2)
    settingWindow.protocol("WM_DELETE_WINDOW", lambda :applySetting())

but_add = tk.Button(frame1, text="确认添加", command=lambda: [additem(), updata_monthwork()]).grid(row=1, column=14,
                                                                                                   padx=5)
but_weekRpt = tk.Button(frame3, text="生成周报", command=lambda: genweekRpt()).grid(row=0, column=0, padx=5, pady=10)
but_monthRpt = tk.Button(frame3, text="生成月报", command=lambda: genmonthRpt()).grid(row=0, column=1, padx=5, pady=10)
but_weekRpt_undo = tk.Button(frame3, text="取消周报标签", command=lambda: genweekRpt_undo("周报")).grid(row=0, column=2,
                                                                                                        padx=5, pady=10)
but_weekRpt_undo = tk.Button(frame3, text="取消月报标签",
                             command=lambda: [genweekRpt_undo("月报"), updata_monthwork()]).grid(row=0, column=3,
                                                                                                 padx=5, pady=10)
but_deleteRecent = tk.Button(frame3, text="删除最近添加", command=lambda: [deleteRecorNot(), updata_monthwork()]).grid(
    row=1, column=1, padx=5, pady=10)
but_delereall = tk.Button(frame3, text="删除所有信息", command=lambda: [deleteAllorNot(), updata_monthwork()]).grid(
    row=1, column=0, padx=5, pady=10)

but_setting = tk.Button(frame3, text="设置", command=lambda: [creatsettingWindow()]).grid(
    row=0, column=4, padx=5, pady=10)





def addbook():
    if (entry_book_pagenum.get().isdigit() or entry_book_pagenum.get() == '') & entry_book_wordnum.get().isdigit():
        print("添加了稿件")
        if entry_book_pagenum.get() == '':
            book_pagenum = 0
        else:
            book_pagenum = entry_book_pagenum.get()
        cmd_add_book = "INSERT INTO book_scores VALUES(NULL,'" + entry_book_name.get() + "',"+ str(book_pagenum) + "," + entry_book_wordnum.get() +",'"+entry_book_info.get()+"',0); "
        cur.execute(cmd_add_book)
        conn.commit()
        book_name_var.set('')
        book_wordnum_var.set('')
        book_pagenum_var.set('')
        book_info_var.set('')
        refreshbookbox()
        # str_show_tmp.set("稿件名称："+entry_name.get()+"，加工类型："+box_type.get()+"，加工页数："+entry_pagenum.get()+"页"+"加工字数："+entry_wordnum.get()+"万字"+" 系数："+entry_xishu.get())
    else:
        messagebox.showerror("错误", "请正确填写所有信息")
        if not entry_book_pagenum.get().isdigit(): book_pagenum_var.set('')
        if not entry_book_wordnum.get().isdigit(): book_wordnum_var.set('')
        return


but_addbook = tk.Button(frame4, text="确认添加", command=lambda: [addbook()]).grid(row=1, column=9,padx=5)


def getmaxid_in_bookscore():
    find_maxid_cmd = "select max(id) from book_scores;"
    cur.execute(find_maxid_cmd)
    maxid = cur.fetchall()[0]
    # print("maxid:", maxid[0])
    if maxid[0] == None:
        return [0, ]
    else:
        return maxid


def deleteLastBook():
    maxid = getmaxid_in_bookscore()
    # print("maxid:")
    # print(maxid)
    delete_recent_cmd = "delete from book_scores where id=(?)"
    cur.execute(delete_recent_cmd, maxid)
    conn.commit()
    refreshbookbox()

def deleteLastBookorNot():
    result = tk.messagebox.askokcancel(title='', message='删除上一条记录？')
    if result:
        deleteLastBook()

but_deletelastbook = tk.Button(frame4, text="删除最近添加", command=lambda: [deleteLastBookorNot()]).grid(row=1, column=10,padx=5)


def deleteAllBook():
    delete_all_cmd = "DELETE FROM book_scores"
    cur.execute(delete_all_cmd)
    conn.commit()
    refreshbookbox()
def deleteAllBookorNot():
    # 弹出对话框
    result = tk.messagebox.askokcancel(title='', message='删除全部稿件？')
    if result:
        deleteAllBook()

but_deleteallbook = tk.Button(frame4, text="删除全部", command=lambda: [deleteAllBookorNot()]).grid(row=1, column=11,padx=5)


updata_monthwork()
refresh()
refreshbookbox()


def closewindow():
    conn.commit()
    cur.close()
    conn.close()
    window.destroy()


window.protocol("WM_DELETE_WINDOW", closewindow)
window.mainloop()
